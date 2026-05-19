import os
import shutil
import unicodedata
import re

from io import BytesIO

def normalizar(txt):

    if not txt:
        return ""

    txt = txt.lower()

    txt = unicodedata.normalize(
        'NFKD',
        txt
    ).encode(
        'ascii',
        'ignore'
    ).decode('utf-8')

    txt = " ".join(txt.split())

    return txt
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import StreamingResponse
from docx import Document
from docxtpl import DocxTemplate
from io import BytesIO
import shutil
from datetime import datetime

app = FastAPI()

# =========================
# LIMPIEZA
# =========================
def limpiar_nombre(nombre):

    eliminar = [
        "SR.",
        "SRA.",
        "DR.",
        "DRA."
    ]

    nombre = nombre.upper()

    for e in eliminar:
        nombre = nombre.replace(e, "")

    return " ".join(nombre.split()).strip()
    
def obtener_tratamiento(cargo):
    cargo = (cargo or "").lower()
    if any(x in cargo for x in ["gerente", "director", "presidente"]):
        return "Doctor"
    return "Señor"
# =====================================
# DETECTAR GENERO
# =====================================
def detectar_genero(nombre):

    nombre = nombre.lower()

    mujeres = [

        "liliana",
        "maria",
        "paula",
        "andrea",
        "carolina",
        "ana",
        "laura",
        "diana",
        "adriana",
        "luisa",
        "catalina",
        "isabel",
        "patricia",
        "solis"
    ]

    hombres = [

        "carlos",
        "juan",
        "pedro",
        "andres",
        "manuel",
        "ivan",
        "jose",
        "luis"
    ]

    primer = nombre.split()[0]

    if primer in mujeres:
        return "f"

    if primer in hombres:
        return "m"

    return "m"


# =====================================
# PRIMER NOMBRE REAL
# =====================================
def obtener_primer_nombre_real(nombre):

    femeninos = [
        "maria","ana","laura","paula","andrea","carolina",
        "diana","luisa","patricia","camila","valentina",
        "anverly","solange","nancy","lina","adriana"
    ]

    masculinos = [
        "juan","carlos","manuel","ivan","andres","luis",
        "felipe","daniel","miguel","jose","eduardo"
    ]

    palabras = nombre.split()

    for p in palabras:

        pl = p.lower().strip()

        if pl in femeninos or pl in masculinos:
            return p.capitalize()

    return palabras[0].capitalize()


# =====================================
# TRATAMIENTO
# =====================================
def obtener_tratamiento(nombre):

    genero = detectar_genero(nombre)

    if genero == "f":
        return "Señora"

    return "Señor"
# =========================
# FECHA
# =========================
def fecha_es():
    meses = [
        "enero","febrero","marzo","abril","mayo","junio",
        "julio","agosto","septiembre","octubre","noviembre","diciembre"
    ]

    hoy = datetime.now()

    dia = str(hoy.day).zfill(2)

    return f"{dia} de {meses[hoy.month-1]} de {hoy.year}"
    
# =========================
# EXTRAER DATOS (FORMATO REAL)
# =========================
def extraer_datos(doc):

    datos = {
        "nombre": "",
        "primer_nombre": "",
        "cargo": "",
        "compañia": "",
        "correo": "",
        "telefono": "",
        "direccion": "",
        "ciudad": "",
    }

    tabla = None

    # =====================================
    # BUSCAR TABLA DEL LEVANTAMIENTO
    # =====================================
    for t in doc.tables:

        texto_tabla = ""
    
        for row in t.rows:
            for cell in row.cells:
                texto_tabla += cell.text.lower() + " "
    
        if (
            "contacto" in texto_tabla
            and "cargo" in texto_tabla
            and (
                "compañía" in texto_tabla
                or "compania" in texto_tabla
            )
        ):
    
            tabla = t
            break

    if tabla is None:
        return datos

    try:

        # =====================================
        # COMPAÑIA
        # fila 4
        # =====================================
        datos["compañia"] = (
            tabla.rows[3].cells[1].text.strip().upper()
        )
        
        # =====================================
        # NOMBRE
        # fila 5
        # =====================================
        if len(tabla.rows) > 4:
            nombre = tabla.rows[4].cells[1].text.strip()
            
            nombre = (
                nombre.upper()
                .replace("SR.", "")
                .replace("SRA.", "")
                .replace("DR.", "")
                .replace("DRA.", "")
                .strip()
            )
            
            datos["nombre"] = nombre
            
        if nombre:
            datos["primer_nombre"] = (
                nombre.split()[0].title()
            )
        
        # =====================================
        # CORREO
        # fila 5 col 3
        # =====================================
        correo = tabla.rows[4].cells[3].text.strip()
        
        if "@" in correo:
            datos["correo"] = correo
        
        # =====================================
        # CARGO
        # fila 6
        # =====================================
        datos["cargo"] = (
            tabla.rows[5].cells[1].text.strip()
        )
        
        # =====================================
        # TELEFONO
        # fila 6 col 3
        # =====================================
        telefono = (
            tabla.rows[5].cells[3].text.strip()
        )
        
        datos["telefono"] = telefono
        
        # =====================================
        # DIRECCION
        # fila 4 col 3
        # =====================================
        datos["direccion"] = (
            tabla.rows[3].cells[3].text.strip()
        )

        # =====================================
        # CIUDAD
        # =====================================
        for t in doc.tables:
        
            for row in t.rows:
        
                try:
        
                    titulo = row.cells[0].text.lower()
        
                    if "ciudad - lugar" in titulo:
        
                        datos["ciudad"] = (
                            row.cells[1].text.strip()
                        )
        
                        break
        
                except:
                    pass
    except Exception as e:

        print("ERROR EXTRACCION:", e)

    return datos
# =========================
# DETECTAR SERVICIO DESDE TABLA X
# =========================
# =========================
# DETECTAR SERVICIO
# =========================
# =========================
# DETECTAR SERVICIO
# =========================
def detectar_servicio(doc):

    texto = obtener_texto_completo(doc)

    print("===================================")
    print("TEXTO ANALIZADO:")
    print(texto)
    print("===================================")

    # =====================================
    # SCORES
    # =====================================
    scores = {

        "eventos": 0,
        "escolta": 0,
        "vigilancia": 0,
        "electronica": 0,
        "monitoreo": 0,
        "confiabilidad": 0
    }

    # =========================================================
    # SEGURIDAD EN EVENTOS
    # SOLO FRASES MUY ESPECIFICAS
    # =========================================================
    keywords_eventos = [

        "seguridad en eventos",
        "seguridad para eventos",
        "logistica de eventos",
        "logistica evento",
        "personal para eventos",
        "control de acceso para eventos",
        "evento corporativo",
        "evento empresarial",
        "seguridad evento",
        "staff de seguridad para evento"
    ]

    for k in keywords_eventos:

        if k in texto:

            scores["eventos"] += 3

            print("EVENTOS MATCH:", k)

    # =========================================================
    # ESCOLTAS
    # =========================================================
    keywords_escolta = [

        "escolta",
        "conductor escolta",
        "escolta motorizado",
        "esquema de seguridad",
        "proteccion ejecutiva",
        "proteccion personal",
        "acompanamiento armado",
        "acompanante de seguridad",
        "driver escolta",
        "escolta vip"
    ]

    for k in keywords_escolta:

        if k in texto:

            scores["escolta"] += 2

            print("ESCOLTA MATCH:", k)

    # =========================================================
    # VIGILANCIA
    # =========================================================
    keywords_vigilancia = [

        "servicio de vigilancia",
        "vigilancia fisica",
        "guarda de seguridad",
        "guardas de seguridad",
        "puesto de vigilancia",
        "servicio armado",
        "sin arma",
        "vigilante",
        "porteria",
        "control de acceso",
        "ronda de vigilancia",
        "medio de comunicacion",
        "supervisor de puesto",
        "seguridad privada"
    ]

    for k in keywords_vigilancia:

        if k in texto:

            scores["vigilancia"] += 2

            print("VIGILANCIA MATCH:", k)

    # =========================================================
    # SEGURIDAD ELECTRONICA
    # =========================================================
    keywords_electronica = [

        "seguridad electronica",
        "cctv",
        "alarmas",
        "camaras",
        "camaras de seguridad",
        "sensor de movimiento",
        "panel de alarma",
        "video vigilancia",
        "control de acceso biometrico",
        "deteccion de intrusion"
    ]

    for k in keywords_electronica:

        if k in texto:

            scores["electronica"] += 2

            print("ELECTRONICA MATCH:", k)

    # =========================================================
    # MONITOREO
    # =========================================================
    keywords_monitoreo = [

        "monitoreo",
        "central de monitoreo",
        "monitoreo remoto",
        "reaccion motorizada",
        "seguimiento satelital",
        "monitoreo gps"
    ]

    for k in keywords_monitoreo:

        if k in texto:

            scores["monitoreo"] += 2

            print("MONITOREO MATCH:", k)

    # =========================================================
    # CONFIABILIDAD
    # =========================================================
    keywords_confiabilidad = [

        "estudio de seguridad",
        "visita domiciliaria",
        "poligrafo",
        "prueba de poligrafo",
        "analisis de confiabilidad",
        "investigacion de antecedentes",
        "validacion de antecedentes",
        "prueba de confiabilidad"
    ]

    for k in keywords_confiabilidad:

        if k in texto:

            scores["confiabilidad"] += 2

            print("CONFIABILIDAD MATCH:", k)

    # =====================================
    # RESULTADOS
    # =====================================
    print("===================================")
    print("SCORES FINALES:")
    print(scores)
    print("===================================")

    # =====================================
    # OBTENER MAYOR SCORE
    # =====================================
    servicio_final = None
    score_mayor = 0

    for servicio, score in scores.items():

        if score > score_mayor:

            score_mayor = score
            servicio_final = servicio

    # =====================================
    # VALIDAR MINIMO
    # =====================================
    if score_mayor <= 0:

        print("NO SE DETECTO SERVICIO")

        return None

    print("===================================")
    print("SERVICIO FINAL:", servicio_final)
    print("SCORE:", score_mayor)
    print("===================================")

    return servicio_final
    
# =========================
# PLANTILLA
# =========================
def seleccionar_plantilla(
    servicio,
    detalle,
    modalidad
):

    print("==== SELECCIONAR PLANTILLA ====")
    print("SERVICIO:", servicio)
    print("DETALLE:", detalle)
    print("MODALIDAD:", modalidad)

    # =====================================
    # SEGURIDAD EN EVENTOS
    # =====================================
    if servicio == "eventos":

        return (
            "plantillas/"
            "seguridad_en_eventos.docx"
        )

    # =====================================
    # ESCOLTAS
    # =====================================
    if servicio == "escolta":

        # MULTIPLE
        if detalle == "multiple":

            return (
                "plantillas/"
                "escolta_mensual.docx"
            )

        # CONDUCTOR
        if detalle == "conductor":

            return (
                "plantillas/"
                "escolta_conductor.docx"
            )

        # MOTORIZADO
        if detalle == "motorizado":

            return (
                "plantillas/"
                "escolta_motorizado.docx"
            )

        # A PIE
        if detalle == "a_pie":

            return (
                "plantillas/"
                "escolta_apie.docx"
            )

        # GENERAL
        return (
            "plantillas/"
            "escolta_mensual.docx"
        )

    # =====================================
    # VIGILANCIA
    # =====================================
    if servicio == "vigilancia":

        if detalle == "vigilancia":

            return (
                "plantillas/"
                "vigilancia_m.docx"
            )

        if detalle == "armada":

            return (
                "plantillas/"
                "vigilancia_armada.docx"
            )

        if detalle == "sin_arma":

            return (
                "plantillas/"
                "vigilancia_sin_arma.docx"
            )

        return (
            "plantillas/"
            "vigilancia_m.docx"
        )

    # =====================================
    # CONFIABILIDAD
    # =====================================
    if servicio == "confiabilidad":

        return (
            "plantillas/"
            "confiabilidad.docx"
        )

    # =====================================
    # ELECTRONICA
    # =====================================
    if servicio == "electronica":

        return (
            "plantillas/"
            "seguridad_electronica.docx"
        )

    # =====================================
    # MONITOREO
    # =====================================
    if servicio == "monitoreo":

        return (
            "plantillas/"
            "monitoreo.docx"
        )

    # =====================================
    # DEFAULT
    # =====================================
    return (
        "plantillas/"
        "default.docx"
    )

def generar_titulos(servicio, detalle):

    titulo_ref = ""
    titulo_servicio = ""

    # =====================================
    # ESCOLTAS
    # =====================================
    if servicio == "escolta":

        # =====================================
        # VARIAS MODALIDADES
        # =====================================
        if detalle == "multiple":

            titulo_ref = (
                "Propuesta Servicio de Escolta"
            )

            titulo_servicio = (
                "SERVICIOS DE ESCOLTA "
                "- DIFERENTES MODALIDADES"
            )

    # =====================================
    # VIGILANCIA
    # =====================================
    elif servicio == "vigilancia":

        # =====================================
        # MIXTA
        # =====================================
        if detalle == "vigilancia_mixta":

            titulo_ref = (
                "Propuesta Servicios de Vigilancia"
            )

            titulo_servicio = (
                "SERVICIOS DE VIGILANCIA "
                "- ARMADA Y SIN ARMA"
            )

        # =====================================
        # ARMADA
        # =====================================
        elif detalle == "armada":

            titulo_ref = (
                "Propuesta Servicio de Vigilancia Armada"
            )

            titulo_servicio = (
                "SERVICIO DE VIGILANCIA ARMADA"
            )

        # =====================================
        # SIN ARMA
        # =====================================
        else:

            titulo_ref = (
                "Propuesta Servicio de Vigilancia "
                "con Medio de Comunicación, Sin Arma"
            )

            titulo_servicio = (
                "SERVICIO DE VIGILANCIA "
                "CON MEDIO DE COMUNICACIÓN SIN ARMA"
            )

    return titulo_ref, titulo_servicio
# =========================
# OBTENER TEXTO COMPLETO
# =========================
def obtener_texto_completo(doc):

    texto = ""

    # =========================
    # PÁRRAFOS
    # =========================
    for p in doc.paragraphs:

        texto += " " + normalizar(p.text)

    # =========================
    # TABLAS
    # =========================
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                texto += " " + normalizar(cell.text)

    return texto
    
def detectar_detalle(doc):

    texto = obtener_texto_completo(doc)

    print("TEXTO DETALLE:")
    print(texto)

    # =====================================
    # DETECTAR SERVICIO
    # =====================================
    servicio = detectar_servicio(doc)

    print("SERVICIO DETECTADO:", servicio)

    # ==================================================
    # VIGILANCIA
    # ==================================================
    if servicio == "vigilancia":

        tiene_armada = (
            "armado" in texto
            or "armada" in texto
        )

        tiene_sin_arma = (
            "sin arma" in texto
            or "medio de comunicacion" in texto
            or "medio de comunicación" in texto
        )

        print("ARMADA:", tiene_armada)
        print("SIN ARMA:", tiene_sin_arma)

        # MIXTA
        if tiene_armada and tiene_sin_arma:

            print("DETALLE: vigilancia_mixta")

            return "vigilancia_mixta"

        # ARMADA
        if tiene_armada:

            print("DETALLE: armada")

            return "armada"

        # SIN ARMA
        if tiene_sin_arma:

            print("DETALLE: sin_arma")

            return "sin_arma"

        print("DETALLE: vigilancia")

        return "vigilancia"

    # ==================================================
    # ESCOLTAS
    # ==================================================
    if servicio == "escolta":

        tiene_motorizado = (

            "motorizado" in texto
            or "moto" in texto
            or "rodamiento" in texto
        )

        tiene_conductor = (

            "conductor escolta" in texto
            or "conductor" in texto
            or "chofer" in texto
            or "driver" in texto
        )

        tiene_apie = (

            "a pie" in texto
            or "acompanante" in texto
            or "acompañante" in texto
        )

        modalidades = []

        if tiene_motorizado:
            modalidades.append("motorizado")

        if tiene_conductor:
            modalidades.append("conductor")

        if tiene_apie:
            modalidades.append("a_pie")

        print("MODALIDADES ESCOLTA:", modalidades)

        # MULTIPLE
        if len(modalidades) >= 2:

            print("DETALLE: multiple")

            return "multiple"

        # UNA
        if len(modalidades) == 1:

            print("DETALLE:", modalidades[0])

            return modalidades[0]

        print("DETALLE: general")

        return "general"

    return None

# =========================
# EXTRAER CONSECUTIVO
# =========================
def extraer_consecutivo(doc):

    texto = obtener_texto_completo(doc)

    print("BUSCANDO CONSECUTIVO EN:")
    print(texto)

    patrones = [

        # 202405111230
        r"\b\d{10,14}\b",

        # FR-GCO-001
        r"\b[A-Z]{2,10}-[A-Z]{2,10}-\d{1,5}\b",

        # consecutivo: XXX
        r"consecutivo[:\s]*([A-Z0-9\-]+)",

        # No. XXX
        r"no\.\s*([A-Z0-9\-]+)"
    ]

    for patron in patrones:

        resultado = re.search(
            patron,
            texto,
            re.IGNORECASE
        )

        if resultado:

            print(
                "CONSECUTIVO ENCONTRADO:",
                resultado.group(0)
            )

            if resultado.groups():
                return resultado.group(1)

            return resultado.group(0)

    return None
# =========================
# DETECTAR MODALIDAD
# =========================
def detectar_modalidad(doc):

    texto = obtener_texto_completo(doc)

    print("TEXTO MODALIDAD:")
    print(texto)

    # =====================================
    # MENSUAL
    # =====================================
    if (
        "mensual" in texto
        or "mes" in texto
    ):
        return "m"

    # =====================================
    # FORTALECIMIENTO
    # =====================================
    if "fortalecimiento" in texto:
        return "f"

    # =====================================
    # EVENTO
    # =====================================
    if "evento" in texto:
        return "e"

    # =====================================
    # DEFAULT
    # =====================================
    return "m"
    
# =========================
# API
# =========================

from fastapi import Form

@app.post("/procesar/")
async def procesar(
    file: UploadFile = File(...),
    servicio_manual: str = Form(None)
):

    try:

        temp = "entrada.docx"

        # =========================
        # GUARDAR ARCHIVO
        # =========================
        with open(temp, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        file.file.close()

        # =========================
        # LEER DOCUMENTO
        # =========================
        doc = Document(temp)

        # =========================
        # EXTRAER DATOS
        # =========================
        datos = extraer_datos(doc)

        print("DATOS EXTRAIDOS:", datos)

        # =========================
        # VALIDAR NOMBRE
        # =========================
        if not datos.get("nombre"):
            datos["nombre"] = "CLIENTE"

        if not datos.get("primer_nombre"):
            datos["primer_nombre"] = "Cliente"

        # =========================
        # DETECTAR SERVICIO
        # =========================
        servicio = detectar_servicio(doc)

        # 🔥 SERVICIO MANUAL
        if servicio_manual and servicio_manual != "string":
            servicio = servicio_manual

        # 🔥 SI NO DETECTA
        if servicio is None:

            return {
                "requiere_seleccion": True,
                "mensaje": "No se pudo detectar el servicio",
                "opciones": [
                    "vigilancia",
                    "escolta",
                    "confiabilidad",
                    "electronica",
                    "monitoreo"
                ]
            }

        print("SERVICIO:", servicio)

        # =========================
        # DETECTAR DETALLE
        # =========================
        detalle = detectar_detalle(doc)
        
        modalidad = detectar_modalidad(doc)
        
        print("SERVICIO FINAL:", servicio)
        print("DETALLE FINAL:", detalle)
        print("MODALIDAD FINAL:", modalidad)

        # =========================
        # DETECTAR MODALIDAD
        # =========================
        modalidad = detectar_modalidad(doc)

        print("MODALIDAD:", modalidad)
        
        print("SERVICIO DETECTADO:", servicio)
        print("DETALLE DETECTADO:", detalle)
        print("MODALIDAD DETECTADA:", modalidad)
        # =========================
        # SELECCIONAR PLANTILLA
        # =========================
        plantilla = seleccionar_plantilla(
            servicio,
            detalle,
            modalidad
        )

        print("PLANTILLA:", plantilla)
        titulo_ref, titulo_servicio = generar_titulos(
            servicio,
            detalle
        )
        print("PLANTILLA FINAL:", plantilla)
        ## =========================
        # REEMPLAZOS
        # =========================
        titulo_ref, titulo_servicio = generar_titulos(
            servicio,
            detalle
        )
        
        # =========================
        # GENERO Y SALUDO
        # =========================
        
        tratamiento = obtener_tratamiento(datos["nombre"])
        
        primer_nombre = obtener_primer_nombre_real(
            datos["nombre"]
        )
        
        genero = detectar_genero(
            datos["nombre"]
        )
        
        saludo_base = "Estimado"
        
        if genero == "f":
            saludo_base = "Estimada"
        # =========================
        # CONSECUTIVO
        # =========================
        
        consecutivo = extraer_consecutivo(doc)
        
        if not consecutivo:
        
            consecutivo = datetime.now().strftime(
                "%Y%m%d%H%M"
            )
        
        print("CONSECUTIVO:", consecutivo)

        # =========================
        # REEMPLAZOS
        # =========================

        
        reemplazos = {

           "consecutivo": consecutivo,
        
            "fecha": fecha_es(),
        
            "tratamiento": tratamiento,
        
            "nombre": datos["nombre"],
        
            "primer_nombre": primer_nombre,
        
            "cargo": datos["cargo"],
        
            "compañia": datos["compañia"],
        
            "correo": datos["correo"],
        
            "telefono": datos["telefono"],
        
            "direccion": datos["direccion"],
        
            "ciudad": datos["ciudad"],
        
            "alcance": datos["ciudad"],
        
            # SALUDO
            "saludo": (
                f"{saludo_base} "
                f"{tratamiento.lower()} "
                f"{primer_nombre}"
            ),
        
            # TITULOS
            "titulo_ref": titulo_ref,
        
            "titulo_servicio": titulo_servicio
        
        }
        # =========================
        # GENERAR DOCUMENTO
        # =========================
        print("PLANTILLA FINAL:", plantilla)
        
        doc_tpl = DocxTemplate(plantilla)

        doc_tpl.render(reemplazos)

        # =========================
        # GUARDAR EN MEMORIA
        # =========================
        buffer = BytesIO()

        doc_tpl.save(buffer)

        buffer.seek(0)

        # =========================
        # RESPUESTA
        # =========================
        return StreamingResponse(

            buffer,

            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),

            headers={
                "Content-Disposition":
                "attachment; filename=resultado.docx"
            }
        )

    except Exception as e:

        print("ERROR:", str(e))

        return {
            "error": str(e)
        }
